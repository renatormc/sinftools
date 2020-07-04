from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
import tempfile
from pathlib import Path
from models import *
from database import db_session
import os
import settings
import jinja2
import report_docx.filters as filters
import settings


class ReportMaker:
    def __init__(self):
        self.tempdir = Path(tempfile.gettempdir())
        self.template_dir = Path(settings.app_dir / "report_docx/templates")
        self.jinja_env = jinja2.Environment()
        for name, function in filters.__dict__.items():
            self.jinja_env.filters[name] = function

    def make_chat(self, query, path=None, chat=None):
        if not path:
            path = self.tempdir / "generated_doc.docx"
        self.tpl = DocxTemplate(self.template_dir / "chat.docx")
        messages = query.all()
        for i, message in enumerate(messages):
            messages[i].data = self._get_data(message)
        context = {'messages': messages, 'chat': chat}
        self.tpl.render(context, self.jinja_env)
        self.tpl.save(path)
        return path

    # def make_images(self, query, n_cols=3):
    #     width = int(150/n_cols)
    #     self.tpl = DocxTemplate(self.template_dir / "images.docx")
    #     files = query.filter(File.type_ == 'image').all()
    #     n = len(files)
    #     n_rows = int(n/n_cols)
    #     if n%n_cols > 0:
    #         n_rows += 1
    #     rows = [[self._get_inline_image(f, width=width) for f in files[i*n_cols: i*n_cols + n_cols]] for i in range(n_rows)]
    #     context = {'rows': rows}
    #     self.tpl.render(context, self.jinja_env)
    #     path = self.tempdir / "generated_doc.docx"
    #     self.tpl.save(path)
    #     return path


    def make_images(self, query, n_cols=3):
        width = Mm(int(150/n_cols))
        files = query.all()
        n = len(files)
        n_rows = int(n/n_cols)
        if n%n_cols > 0:
            n_rows += 1
        rows = [[f for f in files[i*n_cols: i*n_cols + n_cols]] for i in range(n_rows)]
        doc=Document()
        table = doc.add_table(rows = len(rows), cols = n_cols)
        table.style = 'Table Grid'
        path = self.tempdir / "generated_doc.docx"
        for i, row in enumerate(rows):
            for j, file_ in enumerate(row):
                r = table.rows[i].cells[j].paragraphs[0].add_run()
                try:
                    p = file_.path if file_.type_ == 'image' else file_.thumb_path
                    r.add_picture(p, width = width)
                except:
                    table.rows[i].cells[j].text = file_.path
               
        doc.save(path)
        return path

    def make_video_thumbs(self, query, n_cols=3):
        width = int(150/n_cols)
        self.tpl = DocxTemplate(self.template_dir / "images.docx")
        files = query.filter(File.type_ == 'video', File.analise_thumb != None).all()
        n = len(files)
        n_rows = int(n/n_cols)
        if n%n_cols > 0:
            n_rows += 1
        rows = [[self._get_inline_image_thumb(f, width=width) for f in files[i*n_cols: i*n_cols + n_cols]] for i in range(n_rows)]
        context = {'rows': rows}
        self.tpl.render(context, self.jinja_env)
        path = self.tempdir / "generated_doc.docx"
        self.tpl.save(path)
        return path

    # def make_sms(self, query):
    #     doc=Document()
    #     table = doc.add_table(rows = query.count() + 1, cols = 4)
    #     table.style = 'Table Grid'
    #     r = table.rows[i].cells[j].paragraphs[0].add_run()
        
    def _get_inline_image(self, file, width=30):
        path = settings.work_dir / file.path
        if path.exists():
            img = InlineImage(self.tpl, str(path), width=Mm(width))
            if img:
                return img
        return "-- Imagem indisponível --"

    def _get_inline_image_thumb(self, file, width=30):
        path = settings.work_dir / file.thumb_path
        if path.exists():
            img = InlineImage(self.tpl, str(path), width=Mm(width))
            if img:
                return img
        return "-- Imagem indisponível --"


    def _get_data(self, message):
        if message.body:
            return message.body
        if message.attachments.count() > 0:
            attachment = message.attachments[0]
            if attachment.type_ == 'audio':
                return '--Mensagem de áudio--'
            if attachment.type_ == 'image':
                return InlineImage(self.tpl, str(settings.work_dir / attachment.path), width=Mm(30))
            if attachment.type_ == 'video':
                if attachment.has_thumb:
                    return InlineImage(self.tpl, str(settings.work_dir / attachment.thumb_path), width=Mm(40))
                else:
                    return "--Video (miniatura não disponível)--"
            return "--Arquivo--"
        return ""


if __name__ == "__main__":
    rm = ReportMaker()
    rm.make_chat(1)
